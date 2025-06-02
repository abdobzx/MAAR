"""
Document ingestion agent for the Enterprise RAG System.
Handles multi-format document extraction, metadata extraction, and OCR.
"""

import asyncio
import mimetypes
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from uuid import UUID

import aiofiles
import whisper
from docx import Document as DocxDocument
from PIL import Image
from PyPDF2 import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from core.exceptions import DocumentProcessingError, ErrorCodes
from core.logging import LoggerMixin, log_agent_action, log_error
from core.models import Document, DocumentMetadata, DocumentStatus, DocumentType
from database.models import Document as DBDocument


class DocumentExtractor:
    """Base class for document extractors."""
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract content and metadata from document."""
        raise NotImplementedError


class PDFExtractor(DocumentExtractor):
    """PDF document extractor."""
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from PDF."""
        try:
            with open(file_path, 'rb') as file:
                reader = PdfReader(file)
                
                # Extract text
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                # Extract metadata
                info = reader.metadata
                metadata = DocumentMetadata(
                    title=info.get('/Title') if info else None,
                    author=info.get('/Author') if info else None,
                    page_count=len(reader.pages),
                    content_type="application/pdf"
                )
                
                return text.strip(), metadata
                
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract PDF content: {str(e)}",
                error_code=ErrorCodes.EXTRACTION_FAILED
            )


class DOCXExtractor(DocumentExtractor):
    """DOCX document extractor."""
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text and metadata from DOCX."""
        try:
            doc = DocxDocument(file_path)
            
            # Extract text
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Extract metadata
            properties = doc.core_properties
            metadata = DocumentMetadata(
                title=properties.title,
                author=properties.author,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
            return text.strip(), metadata
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract DOCX content: {str(e)}",
                error_code=ErrorCodes.EXTRACTION_FAILED
            )


class TextExtractor(DocumentExtractor):
    """Plain text extractor."""
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                
            metadata = DocumentMetadata(
                content_type="text/plain"
            )
            
            return text.strip(), metadata
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract text content: {str(e)}",
                error_code=ErrorCodes.EXTRACTION_FAILED
            )


class HTMLExtractor(DocumentExtractor):
    """HTML document extractor."""
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from HTML file."""
        try:
            from bs4 import BeautifulSoup
            
            with open(file_path, 'r', encoding='utf-8') as file:
                soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Extract text
            text = soup.get_text(separator='\n', strip=True)
            
            # Extract metadata
            title = soup.find('title')
            title_text = title.get_text() if title else None
            
            metadata = DocumentMetadata(
                title=title_text,
                content_type="text/html"
            )
            
            return text, metadata
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to extract HTML content: {str(e)}",
                error_code=ErrorCodes.EXTRACTION_FAILED
            )


class AudioExtractor(DocumentExtractor):
    """Audio file extractor using Whisper."""
    
    def __init__(self, model_name: str = "base"):
        self.model = whisper.load_model(model_name)
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from audio using Whisper."""
        try:
            result = self.model.transcribe(file_path)
            text = result["text"]
            
            metadata = DocumentMetadata(
                language=result.get("language"),
                content_type=mimetypes.guess_type(file_path)[0] or "audio/*"
            )
            
            return text.strip(), metadata
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to transcribe audio: {str(e)}",
                error_code=ErrorCodes.AUDIO_TRANSCRIPTION_FAILED
            )


class OCRExtractor(DocumentExtractor):
    """OCR extractor for images."""
    
    def __init__(self, language: str = "eng"):
        self.language = language
    
    def extract(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Extract text from image using OCR."""
        try:
            import pytesseract
            
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang=self.language)
            
            metadata = DocumentMetadata(
                content_type=mimetypes.guess_type(file_path)[0] or "image/*"
            )
            
            return text.strip(), metadata
            
        except Exception as e:
            raise DocumentProcessingError(
                f"Failed to perform OCR: {str(e)}",
                error_code=ErrorCodes.OCR_FAILED
            )


class IngestionAgent(LoggerMixin):
    """Document ingestion agent."""
    
    def __init__(self):
        self.extractors = {
            DocumentType.PDF: PDFExtractor(),
            DocumentType.DOCX: DOCXExtractor(),
            DocumentType.TXT: TextExtractor(),
            DocumentType.HTML: HTMLExtractor(),
            DocumentType.AUDIO: AudioExtractor(),
            DocumentType.IMAGE: OCRExtractor(),
        }
    
    def _get_document_type(self, file_path: str) -> DocumentType:
        """Determine document type from file extension."""
        extension = Path(file_path).suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.TXT,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML,
            '.mp3': DocumentType.AUDIO,
            '.wav': DocumentType.AUDIO,
            '.m4a': DocumentType.AUDIO,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.png': DocumentType.IMAGE,
            '.bmp': DocumentType.IMAGE,
            '.tiff': DocumentType.IMAGE,
        }
        
        if extension not in type_mapping:
            raise DocumentProcessingError(
                f"Unsupported file type: {extension}",
                error_code=ErrorCodes.UNSUPPORTED_FORMAT
            )
        
        return type_mapping[extension]
    
    async def process_document(
        self,
        file_path: str,
        original_filename: str,
        user_id: Optional[UUID] = None,
        organization_id: Optional[UUID] = None,
        custom_metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """Process a document and extract content and metadata."""
        
        log_agent_action(
            agent_name="IngestionAgent",
            action="process_document",
            file_path=file_path,
            user_id=str(user_id) if user_id else None
        )
        
        try:
            # Determine document type
            document_type = self._get_document_type(file_path)
            
            # Extract content and metadata
            extractor = self.extractors[document_type]
            content, metadata = extractor.extract(file_path)
            
            # Add custom metadata
            if custom_metadata:
                metadata.custom_fields.update(custom_metadata)
            
            # Get file size
            file_size = Path(file_path).stat().st_size
            metadata.file_size = file_size
            
            # Create document model
            document = Document(
                filename=Path(file_path).name,
                original_filename=original_filename,
                file_path=file_path,
                document_type=document_type,
                status=DocumentStatus.COMPLETED,
                metadata=metadata,
                content=content,
                user_id=user_id,
                organization_id=organization_id
            )
            
            self.logger.info(
                "Document processed successfully",
                document_id=str(document.id),
                document_type=document_type.value,
                content_length=len(content)
            )
            
            return document
            
        except Exception as e:
            log_error(e, {
                "agent": "IngestionAgent",
                "file_path": file_path,
                "user_id": str(user_id) if user_id else None
            })
            
            # Create failed document
            document = Document(
                filename=Path(file_path).name,
                original_filename=original_filename,
                file_path=file_path,
                document_type=self._get_document_type(file_path),
                status=DocumentStatus.FAILED,
                processing_error=str(e),
                user_id=user_id,
                organization_id=organization_id
            )
            
            return document
    
    async def save_document(
        self,
        document: Document,
        db_session: AsyncSession
    ) -> DBDocument:
        """Save document to database."""
        
        try:
            db_document = DBDocument(
                id=document.id,
                filename=document.filename,
                original_filename=document.original_filename,
                file_path=document.file_path,
                document_type=document.document_type,
                status=document.status,
                content=document.content,
                metadata=document.metadata.dict(),
                processing_error=document.processing_error,
                user_id=document.user_id,
                organization_id=document.organization_id
            )
            
            db_session.add(db_document)
            await db_session.commit()
            await db_session.refresh(db_document)
            
            self.logger.info(
                "Document saved to database",
                document_id=str(db_document.id)
            )
            
            return db_document
            
        except Exception as e:
            await db_session.rollback()
            log_error(e, {
                "agent": "IngestionAgent",
                "document_id": str(document.id)
            })
            raise
    
    async def process_batch(
        self,
        file_paths: List[str],
        user_id: Optional[UUID] = None,
        organization_id: Optional[UUID] = None,
        max_workers: int = 4
    ) -> List[Document]:
        """Process multiple documents concurrently."""
        
        semaphore = asyncio.Semaphore(max_workers)
        
        async def process_single(file_path: str) -> Document:
            async with semaphore:
                return await self.process_document(
                    file_path=file_path,
                    original_filename=Path(file_path).name,
                    user_id=user_id,
                    organization_id=organization_id
                )
        
        tasks = [process_single(fp) for fp in file_paths]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        documents = []
        for result in results:
            if isinstance(result, Exception):
                log_error(result, {"agent": "IngestionAgent", "operation": "batch_processing"})
            else:
                documents.append(result)
        
        return documents
