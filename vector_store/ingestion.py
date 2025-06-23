"""
Module d'ingestion et de traitement des documents
Support pour PDF, TXT, DOCX, et autres formats
"""

import os
import uuid
import logging
from typing import List, Dict, Any, Optional, Union, IO
from datetime import datetime
import asyncio
from pathlib import Path

# Imports pour le traitement de documents
try:
    import fitz  # PyMuPDF pour PDF
except ImportError:
    fitz = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import chardet
except ImportError:
    chardet = None

from .models import Document, DocumentChunk, DocumentStatus

logger = logging.getLogger(__name__)


class DocumentIngestion:
    """
    Gestionnaire d'ingestion et de traitement des documents
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialise le système d'ingestion
        
        Args:
            config: Configuration pour l'ingestion
        """
        self.config = config or {}
        
        # Configuration du chunking
        self.chunk_size = self.config.get("chunk_size", 512)
        self.chunk_overlap = self.config.get("chunk_overlap", 50)
        self.min_chunk_size = self.config.get("min_chunk_size", 50)
        
        # Extensions supportées
        self.supported_extensions = {
            '.txt': self._process_text,
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.md': self._process_text,
            '.json': self._process_json,
        }
        
        logger.info("DocumentIngestion initialisé")
    
    async def ingest_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> Document:
        """
        Ingère un fichier et le convertit en Document
        
        Args:
            file_path: Chemin vers le fichier
            metadata: Métadonnées additionnelles
            
        Returns:
            Document traité
        """
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                raise FileNotFoundError(f"Fichier non trouvé: {file_path}")
            
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                raise ValueError(f"Extension non supportée: {extension}")
            
            # Extraire le contenu
            processor = self.supported_extensions[extension]
            content = await processor(file_path)
            
            # Créer le document
            doc_metadata = metadata or {}
            doc_metadata.update({
                "file_path": str(file_path),
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "file_extension": extension,
                "processed_at": datetime.now().isoformat()
            })
            
            document = Document(
                id=str(uuid.uuid4()),
                content=content,
                metadata=doc_metadata,
                source=str(file_path),
                title=file_path.stem,
                status=DocumentStatus.PROCESSING
            )
            
            logger.info(f"Document ingéré: {file_path.name} ({len(content)} caractères)")
            return document
            
        except Exception as e:
            logger.error(f"Erreur lors de l'ingestion de {file_path}: {e}")
            raise
    
    async def ingest_directory(
        self,
        directory_path: Union[str, Path],
        recursive: bool = True,
        metadata: Optional[Dict[str, Any]] = None
    ) -> List[Document]:
        """
        Ingère tous les fichiers d'un répertoire
        
        Args:
            directory_path: Chemin du répertoire
            recursive: Parcours récursif
            metadata: Métadonnées communes
            
        Returns:
            Liste des documents traités
        """
        try:
            directory_path = Path(directory_path)
            
            if not directory_path.exists() or not directory_path.is_dir():
                raise ValueError(f"Répertoire invalide: {directory_path}")
            
            # Trouver tous les fichiers supportés
            pattern = "**/*" if recursive else "*"
            files = []
            
            for file_path in directory_path.glob(pattern):
                if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                    files.append(file_path)
            
            # Traiter les fichiers en parallèle (limité)
            semaphore = asyncio.Semaphore(5)  # Max 5 fichiers en parallèle
            
            async def process_file(file_path):
                async with semaphore:
                    try:
                        return await self.ingest_file(file_path, metadata)
                    except Exception as e:
                        logger.error(f"Échec du traitement de {file_path}: {e}")
                        return None
            
            tasks = [process_file(file_path) for file_path in files]
            results = await asyncio.gather(*tasks, return_exceptions=True)
            
            # Filtrer les résultats valides
            documents = [doc for doc in results if isinstance(doc, Document)]
            
            logger.info(f"Répertoire ingéré: {len(documents)}/{len(files)} fichiers traités")
            return documents
            
        except Exception as e:
            logger.error(f"Erreur lors de l'ingestion du répertoire: {e}")
            raise
    
    async def chunk_document(
        self,
        document: Document,
        custom_chunk_size: Optional[int] = None
    ) -> List[DocumentChunk]:
        """
        Découpe un document en chunks
        
        Args:
            document: Document à découper
            custom_chunk_size: Taille personnalisée des chunks
            
        Returns:
            Liste des chunks créés
        """
        try:
            chunk_size = custom_chunk_size or self.chunk_size
            content = document.content
            
            if len(content) <= chunk_size:
                # Document assez petit, un seul chunk
                chunk = DocumentChunk(
                    id=f"{document.id}_chunk_0",
                    document_id=document.id,
                    content=content,
                    chunk_index=0,
                    start_char=0,
                    end_char=len(content),
                    metadata={
                        **document.metadata,
                        "chunk_type": "single",
                        "original_length": len(content)
                    }
                )
                return [chunk]
            
            # Découpage avec overlap
            chunks = []
            start = 0
            chunk_index = 0
            
            while start < len(content):
                end = min(start + chunk_size, len(content))
                
                # Essayer de couper à un délimiteur naturel
                if end < len(content):
                    # Chercher le dernier délimiteur dans les 100 derniers caractères
                    search_start = max(start, end - 100)
                    search_text = content[search_start:end]
                    
                    for delimiter in ['\n\n', '. ', '.\n', '\n', ' ']:
                        pos = search_text.rfind(delimiter)
                        if pos != -1:
                            end = search_start + pos + len(delimiter)
                            break
                
                chunk_content = content[start:end].strip()
                
                # Vérifier la taille minimale
                if len(chunk_content) >= self.min_chunk_size:
                    chunk = DocumentChunk(
                        id=f"{document.id}_chunk_{chunk_index}",
                        document_id=document.id,
                        content=chunk_content,
                        chunk_index=chunk_index,
                        start_char=start,
                        end_char=end,
                        metadata={
                            **document.metadata,
                            "chunk_type": "split",
                            "original_length": len(content),
                            "chunk_size": len(chunk_content)
                        }
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Avancer avec overlap
                start = end - self.chunk_overlap if end > self.chunk_overlap else end
                
                if start >= len(content):
                    break
            
            logger.debug(f"Document {document.id} découpé en {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Erreur lors du chunking: {e}")
            return []
    
    # Processeurs spécialisés par type de fichier
    
    async def _process_text(self, file_path: Path) -> str:
        """Traite un fichier texte"""
        try:
            # Détecter l'encodage
            with open(file_path, 'rb') as f:
                raw_data = f.read()
            
            encoding = 'utf-8'
            if chardet:
                detected = chardet.detect(raw_data)
                if detected['confidence'] > 0.8:
                    encoding = detected['encoding']
            
            # Lire le contenu
            with open(file_path, 'r', encoding=encoding, errors='ignore') as f:
                content = f.read()
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du fichier texte {file_path}: {e}")
            raise
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Traite un fichier PDF"""
        if fitz is None:
            raise ImportError("PyMuPDF non installé. Installez avec: pip install PyMuPDF")
        
        try:
            doc = fitz.open(file_path)
            content = ""
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                content += f"\n--- Page {page_num + 1} ---\n{text}"
            
            doc.close()
            return content.strip()
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du PDF {file_path}: {e}")
            raise
    
    async def _process_docx(self, file_path: Path) -> str:
        """Traite un fichier DOCX"""
        if DocxDocument is None:
            raise ImportError("python-docx non installé. Installez avec: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Traiter les tableaux
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    content += row_text + "\n"
            
            return content.strip()
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du DOCX {file_path}: {e}")
            raise
    
    async def _process_json(self, file_path: Path) -> str:
        """Traite un fichier JSON"""
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convertir le JSON en texte lisible
            if isinstance(data, dict):
                content = self._dict_to_text(data)
            elif isinstance(data, list):
                content = self._list_to_text(data)
            else:
                content = str(data)
            
            return content
            
        except Exception as e:
            logger.error(f"Erreur lors du traitement du JSON {file_path}: {e}")
            raise
    
    def _dict_to_text(self, data: dict, prefix: str = "") -> str:
        """Convertit un dictionnaire en texte"""
        lines = []
        for key, value in data.items():
            if isinstance(value, dict):
                lines.append(f"{prefix}{key}:")
                lines.append(self._dict_to_text(value, prefix + "  "))
            elif isinstance(value, list):
                lines.append(f"{prefix}{key}:")
                lines.append(self._list_to_text(value, prefix + "  "))
            else:
                lines.append(f"{prefix}{key}: {value}")
        return "\n".join(lines)
    
    def _list_to_text(self, data: list, prefix: str = "") -> str:
        """Convertit une liste en texte"""
        lines = []
        for i, item in enumerate(data):
            if isinstance(item, dict):
                lines.append(f"{prefix}Item {i+1}:")
                lines.append(self._dict_to_text(item, prefix + "  "))
            elif isinstance(item, list):
                lines.append(f"{prefix}Item {i+1}:")
                lines.append(self._list_to_text(item, prefix + "  "))
            else:
                lines.append(f"{prefix}- {item}")
        return "\n".join(lines)
    
    def get_supported_extensions(self) -> List[str]:
        """Retourne la liste des extensions supportées"""
        return list(self.supported_extensions.keys())
    
    def get_stats(self) -> Dict[str, Any]:
        """Retourne les statistiques d'ingestion"""
        return {
            "supported_extensions": self.get_supported_extensions(),
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "min_chunk_size": self.min_chunk_size,
            "pdf_support": fitz is not None,
            "docx_support": DocxDocument is not None,
            "encoding_detection": chardet is not None
        }
